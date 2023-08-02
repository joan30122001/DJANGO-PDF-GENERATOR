from django.shortcuts import render, redirect
from django.views.generic import View
from .forms import TitleForm, SubTitleForm, CoverForm, BoardTitleForm, BoardElementForm
from .models import Title, SubTitle, Cover, BoardTitle, BoardElement
from .forms import RangeTitleForm, RangeSubTitleForm, RangeBoardTitleForm, RangeBoardElementForm
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from xhtml2pdf import pisa
import os
from django.conf import settings
from django.template.loader import get_template
from django.template import Template
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, PageTemplate, Frame, Paragraph
from reportlab.lib.styles import ParagraphStyle
from PyPDF2 import PdfReader, PdfWriter
from django.template import Context
from reportlab.lib.units import inch
from .utils import fetch_resources
from django.db.models import Sum
from django.http import JsonResponse
# from weasyprint import HTML
from django.template.loader import render_to_string
import PyPDF2
from django.template.loader import get_template
from docx import Document
from docx.shared import Inches
from django.http import FileResponse




from django.views.generic import CreateView
from django.shortcuts import render
from django.core.files.storage import FileSystemStorage
from.models import ChartImage




def cover(request):
    if request.method == 'POST':
        form = CoverForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            form = CoverForm()
            # return redirect('input_list')
    else:
        form = CoverForm()
    return render(request, 'cover.html', {'form': form})



def title(request):
    if request.method == 'POST':
        form = TitleForm(request.POST)
        if form.is_valid():
            form.save()
            # return redirect('input_list')
    else:
        form = TitleForm()
    return render(request, 'title.html', {'form': form})



def subtitle(request):
    if request.method == 'POST':
        form = SubTitleForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            # return redirect('input_list')
    else:
        form = SubTitleForm()
    return render(request, 'subtitle.html', {'form': form})



def boardtitle(request):
    if request.method == 'POST':
        form = BoardTitleForm(request.POST)
        if form.is_valid():
            form.save()
            # return redirect('input_list')
    else:
        form = BoardTitleForm()
    return render(request, 'boardtitle.html', {'form': form})



def boardelement(request):
    if request.method == 'POST':
        form = BoardElementForm(request.POST)
        if form.is_valid():
            form.save()
            # return redirect('input_list')
    else:
        form = BoardElementForm()
    return render(request, 'boardelement.html', {'form': form})



def chart(request):
    labels = []
    data = []

    queryset = BoardElement.objects.values('boardtitle__title').annotate(element=Sum('element')).order_by('-element')
    for entry in queryset:
        labels.append(entry['boardtitle__title'])
        data.append(entry['element'])
    
    return JsonResponse(data={
        'labels': labels,
        'data': data,
    })



# def render_to_pdf(template_src, context_dict={}):
#     if template_src is None:
#         html = context_dict['html']
#     else:
#         template = get_template(template_src)
#         html = template.render(context_dict)
#     response = HttpResponse(content_type='application/pdf')
#     response['Content-Disposition'] = 'attachment; filename="preview.pdf"'
#     pisa_status = pisa.CreatePDF(
#         html, dest=response, encoding='utf-8'
#     )
#     if pisa_status.err:
#         return HttpResponse('Failed to generate PDF', status=500)
#     return response



def preview(request):
    # cover = Cover.objects.all()
    # cover = Cover.objects.get(id=cover_id)
    cover = Cover.objects.first()

    titles = Title.objects.all()
    subtitles_by_title = {}
    for title in titles:
        subtitles_by_title[title] = SubTitle.objects.filter(title=title)
    range_title_form = RangeTitleForm()
    range_subTitle_form = RangeSubTitleForm()

    # viens d'être ajouté
    boardtitles = BoardTitle.objects.all()
    boardelements_by_boardtitle = {}
    for boardtitle in boardtitles:
        boardelements_by_boardtitle[boardtitle] = BoardElement.objects.filter(boardtitle=boardtitle)
    range_boardtitle_form = RangeBoardTitleForm()
    range_boardelement_form = RangeBoardElementForm()
    # fin de l'ajout

    context = {
        'subtitles_by_title': subtitles_by_title,
        'range_title_form': range_title_form,
        'range_subTitle_form': range_subTitle_form,
        'cover': cover,

        'boardelements_by_boardtitle': boardelements_by_boardtitle,
        'range_boardtitle_form': range_boardtitle_form,
        'range_boardelement_form': range_boardelement_form,

        # 'boardtitles_by_subtitle': boardtitles_by_subtitle,
    }
    return render(request, 'preview.html', context)



def preview1(request):
    # cover = Cover.objects.all()
    # cover = Cover.objects.get(id=cover_id)
    cover = Cover.objects.first()

    titles = Title.objects.all()
    subtitles_by_title = {}
    for title in titles:
        subtitles_by_title[title] = SubTitle.objects.filter(title=title)
    range_title_form = RangeTitleForm()
    range_subTitle_form = RangeSubTitleForm()

    # viens d'être ajouté
    boardtitles = BoardTitle.objects.all()
    boardelements_by_boardtitle = {}
    for boardtitle in boardtitles:
        boardelements_by_boardtitle[boardtitle] = BoardElement.objects.filter(boardtitle=boardtitle)
    range_boardtitle_form = RangeBoardTitleForm()
    range_boardelement_form = RangeBoardElementForm()
    # fin de l'ajout

    context = {
        'subtitles_by_title': subtitles_by_title,
        'range_title_form': range_title_form,
        'range_subTitle_form': range_subTitle_form,
        'cover': cover,

        'boardelements_by_boardtitle': boardelements_by_boardtitle,
        'range_boardtitle_form': range_boardtitle_form,
        'range_boardelement_form': range_boardelement_form,

        # 'boardtitles_by_subtitle': boardtitles_by_subtitle,
    }
    return render(request, 'preview1.html', context)



# def download_pdf(request):
#     # cover = Cover.objects.all()

    # cover = Cover.objects.first()

    # titles = Title.objects.all()
    # subtitles_by_title = {}
    # for title in titles:
    #     subtitles_by_title[title] = SubTitle.objects.filter(title=title)
    # range_title_form = RangeTitleForm()
    # range_subTitle_form = RangeSubTitleForm()
    # context = {
    #     'subtitles_by_title': subtitles_by_title,
    #     'range_title_form': range_title_form,
    #     'range_subTitle_form': range_subTitle_form,
    #     'cover': cover,
    # }
#     pdf = render_to_pdf('pdf.html', context)
#     return pdf


    # Render only the desired div
    # html = render(request, 'preview.html', context)
    # div_id = 'to_generate'  # Replace with the ID of the desired div
    # start_tag = '<div id="{}">'.format(div_id)
    # end_tag = '</div>'
    # start_index = html.content.decode().find(start_tag)
    # end_index = html.content.decode().find(end_tag, start_index) + len(end_tag)
    # div_html = html.content.decode()[start_index:end_index]
    # # Generate the PDF from the div HTML
    # pdf = render_to_pdf(None, {'html': div_html})
    # return pdf



# class ChartImageUploadView(CreateView):
#     model = ChartImage
#     fields = ['image']
#     success_url = '/chart-image-upload/'

#     def form_valid(self, form):
#         response = super().form_valid(form)
#         # Generate chart image
#         chart_image_url = generate_chart_image()
#         # Save chart image URL to database
#         self.object.image.url = chart_image_url
#         self.object.save()
#         return response

# def generate_chart_image():
#     # Get data from database
#     data = BoardElement.objects.values('boardtitle__title', 'element').annotate(count=models.Count('element'))

#     # Extract labels and values from data
#     labels = [d['boardtitle__title'] for d in data]
#     values = [d['element'] for d in data]

#     # Generate chart image
#     plt.bar(labels, values)
#     plt.xticks(rotation=45)
#     buffer = BytesIO()
#     plt.savefig(buffer, format='png')
#     buffer.seek(0)
#     chart_image = ContentFile(buffer.getvalue())

#     # Save chart image to database
#     chart_image_model = ChartImage.objects.create(image=chart_image)

#     return chart_image_model.image.url



# def generate_pdf(request):
#     template_path = 'pdf.html'
#     cover = Cover.objects.first()
#     titles = Title.objects.all()
#     subtitles_by_title = {}
#     for title in titles:
#         subtitles_by_title[title] = SubTitle.objects.filter(title=title)
#     range_title_form = RangeTitleForm()
#     range_subTitle_form = RangeSubTitleForm()



#     # viens d'être ajouté
#     boardtitles = BoardTitle.objects.all()
#     boardelements_by_boardtitle = {}
#     for boardtitle in boardtitles:
#         boardelements_by_boardtitle[boardtitle] = BoardElement.objects.filter(boardtitle=boardtitle)
#     range_boardtitle_form = RangeBoardTitleForm()
#     range_boardelement_form = RangeBoardElementForm()

#     context = {
#         'subtitles_by_title': subtitles_by_title,
#         'range_title_form': range_title_form,
#         'range_subTitle_form': range_subTitle_form,
#         'cover': cover,

#         'boardelements_by_boardtitle': boardelements_by_boardtitle,
#         'range_boardtitle_form': range_boardtitle_form,
#         'range_boardelement_form': range_boardelement_form,

#     }
#     response = HttpResponse(content_type='application/pdf')
#     response['Content-Disposition'] = 'filename="preview.pdf"'
#     template = get_template(template_path)
#     html = template.render(context)
#     result = BytesIO()
#     pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)
#     # pdf = pisa.CreatePDF(html, dest=result, link_callback=fetch_resources, canvasmaker=page_number)
#     if not pdf.err:
#         response.write(result.getvalue())
#         return response
#     return HttpResponse('Error generating PDF')
    



def template(request):
    return render(request, 'template.html')



# def chart(request):
#     labels = []
#     data = []

#     queryset = BoardElement.objects.values('boardtitle__title').annotate(element=Sum('element')).order_by('-element')
#     for entry in queryset:
#         labels.append(entry['boardtitle__title'])
#         data.append(entry['element'])
    
#     return JsonResponse(data={
#         'labels': labels,
#         'data': data,
#     })



import threading
import matplotlib
matplotlib.use('Agg')
import io
import matplotlib.pyplot as plt
from django.core.files.base import ContentFile
from django.db import models
import queue
from django.urls import reverse


queue = queue.Queue()

def generate_chart_image(queue):
    # Get data from database
    # data = BoardElement.objects.values('boardtitle__title', 'element').annotate(count=models.Count('id')).order_by('element')
    data = BoardElement.objects.values('boardtitle__title').annotate(element=Sum('element')).order_by('-element')

    # Extract labels and values from data
    labels = [d['boardtitle__title'] for d in data]
    values = [d['element'] for d in data]

    # Set the figure size to ensure a visible scale
    plt.figure(figsize=(10, 6))

    # Generate chart image
    plt.bar(labels, values)
    plt.xticks(rotation=0)

    # # Set axis labels and font size
    # plt.xlabel('Board Titles', fontsize=8)
    # plt.ylabel('Element Count', fontsize=12)

    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    chart_image = ContentFile(buffer.getvalue())

    # Put chart image in queue
    queue.put(chart_image)



def charts(request):
    # Create a queue to pass the chart image from the worker thread to the main thread
    # queue = queue.Queue()

    # Generate chart image in a separate thread
    t = threading.Thread(target=generate_chart_image, args=(queue,))
    t.start()

    # Wait for the thread to finish and get the chart image from the queue
    t.join()
    chart_image = queue.get()

    # Return chart image as response
    return HttpResponse(chart_image, content_type='image/png')






def generate_word(request):
    cover = Cover.objects.first()
    titles = Title.objects.all()
    subtitles_by_title = {}
    for title in titles:
        subtitles_by_title[title] = SubTitle.objects.filter(title=title)
    range_title_form = RangeTitleForm()
    range_subTitle_form = RangeSubTitleForm()

    # viens d'être ajouté
    boardtitles = BoardTitle.objects.all()
    boardelements_by_boardtitle = {}
    for boardtitle in boardtitles:
        boardelements_by_boardtitle[boardtitle] = BoardElement.objects.filter(boardtitle=boardtitle)
    range_boardtitle_form = RangeBoardTitleForm()
    range_boardelement_form = RangeBoardElementForm()

    context = {
        'subtitles_by_title': subtitles_by_title,
        'range_title_form': range_title_form,
        'range_subTitle_form': range_subTitle_form,
        'cover': cover,

        'boardelements_by_boardtitle': boardelements_by_boardtitle,
        'range_boardtitle_form': range_boardtitle_form,
        'range_boardelement_form': range_boardelement_form,

    }
    template = get_template('pdf.html')
    # context = {'html_documents': html_documents}
    # return HttpResponse(template.render(context, request))
    rendered_template = template.render(context, request)
    # document = Document()
    # document.add_paragraph(rendered_template)
    # file_name = 'generated_document.docx'
    # document.save(file_name)
    # with open(file_name, 'rb') as file:
    #     response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    #     response['Content-Disposition'] = f'attachment; filename="{file_name}"'
    # os.remove(file_name)
    # return response
    file_path = 'generated_document.docx'
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(rendered_template)
    file.close()
    # return FileResponse(open(file_path, 'rb'), as_attachment=True, filename='converted_document.docx')
    return FileResponse(open(file_path, 'rb'), as_attachment=True, filename='converted_document.docx', content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


















def generate_pdf(request):
    template_path = 'pdf.html'
    cover = Cover.objects.first()
    titles = Title.objects.all()
    subtitles_by_title = {}
    for title in titles:
        subtitles_by_title[title] = SubTitle.objects.filter(title=title)
    range_title_form = RangeTitleForm()
    range_subTitle_form = RangeSubTitleForm()



    # viens d'être ajouté
    boardtitles = BoardTitle.objects.all()
    boardelements_by_boardtitle = {}
    for boardtitle in boardtitles:
        boardelements_by_boardtitle[boardtitle] = BoardElement.objects.filter(boardtitle=boardtitle)
    range_boardtitle_form = RangeBoardTitleForm()
    range_boardelement_form = RangeBoardElementForm()

    #new
    image_url = request.build_absolute_uri(reverse('app:charts'))
    #end new

    context = {
        'subtitles_by_title': subtitles_by_title,
        'range_title_form': range_title_form,
        'range_subTitle_form': range_subTitle_form,
        'cover': cover,

        #new
        'image_url': image_url,
        #end new

        'boardelements_by_boardtitle': boardelements_by_boardtitle,
        'range_boardtitle_form': range_boardtitle_form,
        'range_boardelement_form': range_boardelement_form,

    }
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'filename="preview.pdf"'
    template = get_template(template_path)
    html = template.render(context)
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)
    # pdf = pisa.CreatePDF(html, dest=result, link_callback=fetch_resources, canvasmaker=page_number)
    if not pdf.err:
        response.write(result.getvalue())
        return response
    return HttpResponse('Error generating PDF')